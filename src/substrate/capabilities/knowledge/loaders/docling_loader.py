"""Docling document loader — structure-aware extraction with rich file metadata.

Docling handles layout, tables, code blocks, and images far better than
plain PDF parsers.  This loader pairs it with a metadata pre-extraction
step (pypdf / python-docx / file stats) so that properties like author,
creation date, and title are preserved in every Document's metadata dict.

Three output formats are supported:

* ``"blocks"``   — typed ``ContentBlock`` list (TextBlock, DataBlock, CodeBlock,
                    ImageBlock) that preserves structure for multimodal RAG.
* ``"markdown"`` — Docling's native ``export_to_markdown()`` output, ideal for
                    passing clean, LLM-ready text to embedding models or prompts.
* ``"html"``     — Docling's native ``export_to_html()`` output, useful when the
                    downstream consumer needs rendered structure.

Supported input formats (via Docling): PDF, DOCX, PPTX, HTML, Markdown,
AsciiDoc, images (with OCR).

Requires:  uv add docling

Optional (richer file metadata):
  uv add python-docx   # DOCX/DOC author, dates, etc.
  pypdf is already in project deps for PDF property extraction.
"""

from __future__ import annotations

import hashlib
import io
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Literal, Union

from substrate.capabilities.knowledge.loaders.base import BaseDocumentLoader
from substrate.kernel.core.content import CodeBlock, DataBlock, ImageBlock, TextBlock
from substrate.kernel.storage.vector import Document
from substrate.logger import setup_logging

logger = setup_logging("substrate.capabilities.knowledge.loaders.docling")

# ---------------------------------------------------------------------------
# Docling node-type → ContentBlock mapping
# ---------------------------------------------------------------------------

_HEADING_LABELS = {"section_header", "page_header", "title"}
_CODE_LABELS = {"code"}
_LIST_LABELS = {"list_item"}
_TABLE_LABELS = {"table"}
_IMAGE_LABELS = {"picture", "figure"}
_SKIP_LABELS = {"page_footer", "page_number", "caption"}

OutputFormat = Literal["blocks", "markdown", "html"]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _chunk_id(source: str, index: int, text: str) -> str:
    """Deterministic, content-addressed chunk ID for idempotent re-ingestion."""
    digest = hashlib.sha256(text.encode()).hexdigest()
    return str(uuid.uuid5(uuid.NAMESPACE_URL, f"{source}|{index}|{digest}"))


def _extract_pdf_metadata(path: Path) -> dict[str, Any]:
    """Extract author, title, creation/modification dates from PDF properties."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(path)
        info = reader.metadata or {}

        def _parse_pdf_date(raw: str | None) -> str | None:
            if not raw:
                return None
            raw = raw.lstrip("D:")
            try:
                return datetime.strptime(raw[:14], "%Y%m%d%H%M%S").isoformat()
            except (ValueError, IndexError):
                return raw

        return {
            k: v
            for k, v in {
                "title": info.get("/Title"),
                "author": info.get("/Author"),
                "subject": info.get("/Subject"),
                "creator": info.get("/Creator"),
                "producer": info.get("/Producer"),
                "created_at": _parse_pdf_date(info.get("/CreationDate")),
                "modified_at": _parse_pdf_date(info.get("/ModDate")),
                "page_count": len(reader.pages),
            }.items()
            if v is not None
        }
    except Exception as exc:
        logger.debug("PDF metadata extraction failed for %s: %s", path, exc)
        return {}


def _extract_docx_metadata(path: Path) -> dict[str, Any]:
    """Extract core properties from DOCX/DOC files."""
    try:
        from docx import Document as DocxDoc  # type: ignore[import-unresolved]

        props = DocxDoc(str(path)).core_properties
        return {
            k: v
            for k, v in {
                "title": props.title or None,
                "author": props.author or None,
                "subject": props.subject or None,
                "description": props.comments or None,
                "keywords": props.keywords or None,
                "created_at": props.created.isoformat() if props.created else None,
                "modified_at": props.modified.isoformat() if props.modified else None,
                "last_modified_by": props.last_modified_by or None,
                "revision": props.revision,
            }.items()
            if v is not None
        }
    except ImportError:
        logger.debug("python-docx not installed; DOCX metadata not extracted")
        return {}
    except Exception as exc:
        logger.debug("DOCX metadata extraction failed for %s: %s", path, exc)
        return {}


_FILE_METADATA_EXTRACTORS: dict[str, Any] = {
    ".pdf": _extract_pdf_metadata,
    ".docx": _extract_docx_metadata,
    ".doc": _extract_docx_metadata,
}


def _file_base_metadata(path: Path) -> dict[str, Any]:
    """Always-available metadata from the filesystem."""
    stat = path.stat()
    return {
        "source": str(path),
        "file_name": path.name,
        "file_extension": path.suffix.lower(),
        "file_size_bytes": stat.st_size,
        "fs_modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }


# ---------------------------------------------------------------------------
# DoclingLoader
# ---------------------------------------------------------------------------


class DoclingLoader(BaseDocumentLoader):
    """Structure-aware document loader backed by Docling.

    Args:
        output_format: Controls what ``Document.content`` contains.

            ``"blocks"`` (default)
                Typed ``ContentBlock`` list — ``TextBlock`` for prose/headings,
                ``DataBlock`` for tables, ``CodeBlock`` for fenced code,
                ``ImageBlock`` for figures (if ``extract_images=True``).
                Best for multimodal RAG pipelines.

            ``"markdown"``
                Docling's native ``export_to_markdown()`` output as a single
                ``TextBlock`` per chunk.  Clean, LLM-ready text that preserves
                table formatting and heading hierarchy without extra parsing.
                Best for standard embedding + retrieval pipelines.

            ``"html"``
                Docling's native ``export_to_html()`` as a single ``TextBlock``
                per document.  Useful when the downstream renderer needs full
                HTML structure.  Always produces one Document per file
                (page/section splitting doesn't apply to HTML output).

        chunk_by_page: Produce one ``Document`` per page (``True``, default)
            or one per top-level section / heading (``False``).
            Ignored when ``output_format="html"`` (always one Document).

        extract_images: Include ``ImageBlock`` entries for figures.
            Only applies when ``output_format="blocks"``.
            Off by default — base64 images inflate embedding payloads.

        ocr_enabled: Enable Docling's OCR pass for scanned pages.
            Adds significant latency.
    """

    def __init__(
        self,
        *,
        output_format: OutputFormat = "blocks",
        chunk_by_page: bool = True,
        extract_images: bool = False,
        ocr_enabled: bool = False,
    ) -> None:
        self.output_format = output_format
        self.chunk_by_page = chunk_by_page
        self.extract_images = extract_images
        self.ocr_enabled = ocr_enabled

    # ------------------------------------------------------------------
    # BaseDocumentLoader
    # ------------------------------------------------------------------

    async def load(
        self,
        source: Union[str, Path, bytes],
        *,
        metadata: dict[str, Any] | None = None,
    ) -> list[Document]:
        base_meta = dict(metadata or {})

        # 1. Resolve source + extract file-level metadata before Docling
        if isinstance(source, bytes):
            path: Path | None = None
            source_label = str(base_meta.get("source", "<bytes>"))
        else:
            path = Path(source)
            base_meta.update(_file_base_metadata(path))
            extractor = _FILE_METADATA_EXTRACTORS.get(path.suffix.lower())
            if extractor:
                base_meta.update(extractor(path))
            source_label = str(path)

        base_meta["output_format"] = self.output_format

        # 2. Convert via Docling
        result = self._convert(source, path)

        # 3. Dispatch to the right renderer
        if self.output_format == "html":
            return self._export_html(result, base_meta, source_label)
        if self.output_format == "markdown":
            return self._export_markdown(result, base_meta, source_label)
        # "blocks"
        if self.chunk_by_page:
            return self._chunk_by_page(result, base_meta, source_label)
        return self._chunk_by_section(result, base_meta, source_label)

    # ------------------------------------------------------------------
    # Docling conversion
    # ------------------------------------------------------------------

    def _converter(self) -> Any:
        from docling.document_converter import DocumentConverter, PdfFormatOption  # type: ignore[import-unresolved]
        from docling.datamodel.pipeline_options import PdfPipelineOptions  # type: ignore[import-unresolved]
        from docling.datamodel.base_models import InputFormat  # type: ignore[import-unresolved]

        pipeline_opts = PdfPipelineOptions(do_ocr=self.ocr_enabled)
        return DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_opts)
            }
        )

    def _convert(self, source: Union[str, Path, bytes], path: Path | None) -> Any:
        converter = self._converter()
        if isinstance(source, bytes):
            import os
            import tempfile

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(source)
                tmp_path = tmp.name
            try:
                return converter.convert(tmp_path)
            finally:
                os.unlink(tmp_path)
        return converter.convert(str(path))

    # ------------------------------------------------------------------
    # output_format="markdown"
    # ------------------------------------------------------------------

    def _export_markdown(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        """One Document per page (or section) rendered as Docling markdown."""
        if self.chunk_by_page:
            return self._markdown_by_page(result, base_meta, source_label)
        return self._markdown_by_section(result, base_meta, source_label)

    def _markdown_by_page(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        pages: dict[int, list[Any]] = {}
        for item, _level in result.document.iterate_items():
            if getattr(item, "label", "") in _SKIP_LABELS:
                continue
            prov = getattr(item, "prov", None)
            page_no = prov[0].page_no if prov else 0
            pages.setdefault(page_no, []).append(item)

        total_pages = getattr(result.document, "num_pages", len(pages)) or len(pages)
        docs: list[Document] = []
        for page_no, items in sorted(pages.items()):
            md = self._items_to_markdown(items)
            if not md.strip():
                continue
            meta = {**base_meta, "page_number": page_no, "total_pages": total_pages}
            docs.append(
                Document(
                    content=[TextBlock(text=md)],
                    metadata=meta,
                    id=_chunk_id(source_label, page_no, md),
                )
            )
        return docs

    def _markdown_by_section(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        # Use Docling's whole-document export then split on headings
        full_md: str = result.document.export_to_markdown()
        sections = _split_markdown_by_heading(full_md)
        docs: list[Document] = []
        for idx, (heading, body) in enumerate(sections):
            md = f"## {heading}\n\n{body}".strip() if heading else body.strip()
            if not md:
                continue
            meta = {**base_meta, "section_heading": heading, "section_index": idx}
            docs.append(
                Document(
                    content=[TextBlock(text=md)],
                    metadata=meta,
                    id=_chunk_id(source_label, idx, md),
                )
            )
        return docs

    def _items_to_markdown(self, items: list[Any]) -> str:
        """Render a list of Docling items as a markdown string."""
        parts: list[str] = []
        for item in items:
            label = getattr(item, "label", "")
            if label in _SKIP_LABELS:
                continue
            if label in _TABLE_LABELS:
                try:
                    md = (
                        item.export_to_markdown()
                        if hasattr(item, "export_to_markdown")
                        else ""
                    )
                    if md:
                        parts.append(md)
                        continue
                except Exception:
                    pass
                text = getattr(item, "text", "") or ""
                if text:
                    parts.append(text)
            elif label in _HEADING_LABELS:
                text = (getattr(item, "text", "") or "").strip()
                if text:
                    parts.append(f"## {text}")
            elif label in _CODE_LABELS:
                text = (getattr(item, "text", "") or "").strip()
                if text:
                    parts.append(f"```\n{text}\n```")
            elif label in _LIST_LABELS:
                text = (getattr(item, "text", "") or "").strip()
                if text:
                    parts.append(f"- {text}")
            elif label not in _IMAGE_LABELS:
                text = (getattr(item, "text", "") or "").strip()
                if text:
                    parts.append(text)
        return "\n\n".join(parts)

    # ------------------------------------------------------------------
    # output_format="html"
    # ------------------------------------------------------------------

    def _export_html(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        """Single Document containing the full HTML export."""
        html: str = result.document.export_to_html()
        if not html.strip():
            return []
        return [
            Document(
                content=[TextBlock(text=html)],
                metadata=base_meta,
                id=_chunk_id(source_label, 0, html),
            )
        ]

    # ------------------------------------------------------------------
    # output_format="blocks" — page chunking
    # ------------------------------------------------------------------

    def _chunk_by_page(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        pages: dict[int, list[Any]] = {}
        for item, _level in result.document.iterate_items():
            if getattr(item, "label", "") in _SKIP_LABELS:
                continue
            prov = getattr(item, "prov", None)
            page_no = prov[0].page_no if prov else 0
            pages.setdefault(page_no, []).append(item)

        total_pages = getattr(result.document, "num_pages", len(pages)) or len(pages)
        docs: list[Document] = []
        for page_no, items in sorted(pages.items()):
            blocks = self._items_to_blocks(items)
            if not blocks:
                continue
            text_repr = "\n\n".join(
                b.to_text_repr() for b in blocks if hasattr(b, "to_text_repr")
            )
            meta = {**base_meta, "page_number": page_no, "total_pages": total_pages}
            docs.append(
                Document(
                    content=blocks,
                    metadata=meta,
                    id=_chunk_id(source_label, page_no, text_repr),
                )
            )
        return docs

    # ------------------------------------------------------------------
    # output_format="blocks" — section chunking
    # ------------------------------------------------------------------

    def _chunk_by_section(
        self,
        result: Any,
        base_meta: dict[str, Any],
        source_label: str,
    ) -> list[Document]:
        sections: list[tuple[str, list[Any]]] = []
        current_heading = ""
        current_items: list[Any] = []

        for item, _level in result.document.iterate_items():
            label = getattr(item, "label", "")
            if label in _SKIP_LABELS:
                continue
            if label in _HEADING_LABELS:
                if current_items:
                    sections.append((current_heading, current_items))
                current_items = []
                current_heading = (getattr(item, "text", "") or "").strip()
            else:
                current_items.append(item)

        if current_items:
            sections.append((current_heading, current_items))

        docs: list[Document] = []
        for idx, (heading, items) in enumerate(sections):
            blocks = self._items_to_blocks(items)
            if not blocks:
                continue
            text_repr = "\n\n".join(
                b.to_text_repr() for b in blocks if hasattr(b, "to_text_repr")
            )
            meta = {**base_meta, "section_heading": heading, "section_index": idx}
            docs.append(
                Document(
                    content=blocks,
                    metadata=meta,
                    id=_chunk_id(source_label, idx, text_repr),
                )
            )
        return docs

    # ------------------------------------------------------------------
    # Item → ContentBlock conversion (blocks mode)
    # ------------------------------------------------------------------

    def _items_to_blocks(self, items: list[Any]) -> list[Any]:
        blocks: list[Any] = []
        for item in items:
            block = self._item_to_block(item)
            if block is not None:
                blocks.append(block)
        return blocks

    def _item_to_block(self, item: Any) -> Any:
        label = getattr(item, "label", "")

        if label in _TABLE_LABELS:
            return self._table_to_data_block(item)

        if label in _CODE_LABELS:
            text = (getattr(item, "text", "") or "").strip()
            return CodeBlock(code=text, language="") if text else None

        if label in _IMAGE_LABELS:
            return self._picture_to_image_block(item) if self.extract_images else None

        if label in _LIST_LABELS:
            text = (getattr(item, "text", "") or "").strip()
            return TextBlock(text=f"• {text}") if text else None

        text = (getattr(item, "text", "") or "").strip()
        if not text:
            return None
        if label in _HEADING_LABELS:
            return TextBlock(text=f"## {text}")
        return TextBlock(text=text)

    def _table_to_data_block(self, item: Any) -> DataBlock | TextBlock | None:
        try:
            table_data = getattr(item, "data", None)
            if table_data is None:
                raise AttributeError("no table data")

            grid = getattr(table_data, "grid", None)
            if grid is None:
                md = (
                    item.export_to_markdown()
                    if hasattr(item, "export_to_markdown")
                    else ""
                )
                return TextBlock(text=md) if md else None

            rows: list[list[str]] = []
            for row in grid:
                rows.append(
                    [cell.text if hasattr(cell, "text") else str(cell) for cell in row]
                )
            return DataBlock(data={"rows": rows})
        except Exception as exc:
            logger.debug("Table conversion failed: %s", exc)
            text = (getattr(item, "text", "") or "").strip()
            return TextBlock(text=text) if text else None  # type: ignore[return-value]

    def _picture_to_image_block(self, item: Any) -> ImageBlock | None:
        try:
            img = None
            if hasattr(item, "image") and item.image is not None:
                img = item.image
            elif hasattr(item, "get_image"):
                img = item.get_image()

            if img is None:
                return None

            buf = io.BytesIO()
            if hasattr(img, "save"):
                img.save(buf, format="PNG")
            elif isinstance(img, bytes):
                buf.write(img)
            else:
                return None

            return ImageBlock(data=buf.getvalue(), media_type="image/png")
        except Exception as exc:
            logger.debug("Image extraction failed: %s", exc)
            return None


# ---------------------------------------------------------------------------
# Markdown splitting helper
# ---------------------------------------------------------------------------


def _split_markdown_by_heading(md: str) -> list[tuple[str, str]]:
    """Split a markdown string into (heading, body) pairs on ``##`` headings."""
    sections: list[tuple[str, str]] = []
    current_heading = ""
    current_lines: list[str] = []

    for line in md.splitlines():
        if line.startswith("## "):
            if current_lines or current_heading:
                sections.append((current_heading, "\n".join(current_lines).strip()))
            current_heading = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines or current_heading:
        sections.append((current_heading, "\n".join(current_lines).strip()))

    return sections or [("", md)]
